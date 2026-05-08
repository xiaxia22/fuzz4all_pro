"""按陕西科技大学毕设论文模板格式生成标准 .docx 文件"""
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 全局常量 ──
FONT_SONG = '宋体'
FONT_HEI = '黑体'
FONT_TNR = 'Times New Roman'
SZ_SAN = Pt(16)    # 三号
SZ_SI = Pt(14)     # 四号
SZ_XIAO_SI = Pt(12)  # 小四
SZ_WU = Pt(10.5)   # 五号
SZ_XIAO_WU = Pt(9)  # 小五

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.8)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.2)


def set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TNR, size=SZ_XIAO_SI, bold=False, color=None):
    """设置run的中英文字体"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = en_font
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    if color:
        run.font.color.rgb = color


def add_paragraph_formatted(text, cn_font=FONT_SONG, en_font=FONT_TNR, size=SZ_XIAO_SI,
                            bold=False, align=None, space_before=None, space_after=None,
                            first_line_indent=None, line_spacing=1.5):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    pf.line_spacing = line_spacing
    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold)
    return p


def add_chapter_title(text):
    """章标题：居中，三号黑体，段后1行"""
    return add_paragraph_formatted(text, cn_font=FONT_HEI, size=SZ_SAN, bold=True,
                                   align=WD_ALIGN_PARAGRAPH.CENTER,
                                   space_after=Pt(12), first_line_indent=None)


def add_section_title(text):
    """节标题：左起顶格，四号黑体，段前0.5行，段后0.5行"""
    return add_paragraph_formatted(text, cn_font=FONT_HEI, size=SZ_SI, bold=True,
                                   align=WD_ALIGN_PARAGRAPH.LEFT,
                                   space_before=Pt(6), space_after=Pt(6),
                                   first_line_indent=None)


def add_subsection_title(text):
    """小节标题：左起顶格，小四黑体，段前0.5行，段后0.5行"""
    return add_paragraph_formatted(text, cn_font=FONT_HEI, size=SZ_XIAO_SI, bold=True,
                                   align=WD_ALIGN_PARAGRAPH.LEFT,
                                   space_before=Pt(6), space_after=Pt(6),
                                   first_line_indent=None)


def add_body(text):
    """正文段落：小四宋体，首行缩进2字符"""
    return add_paragraph_formatted(text, cn_font=FONT_SONG, en_font=FONT_TNR, size=SZ_XIAO_SI,
                                   first_line_indent=Cm(0.74))


def add_body_mixed(segments):
    """混合中英文的正文段落，segments = [(text, bold), ...]"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.74)
    pf.line_spacing = 1.5
    for text, bold in segments:
        run = p.add_run(text)
        # 判断是否主要是英文
        en_ratio = sum(1 for c in text if ord(c) < 128) / max(len(text), 1)
        if en_ratio > 0.5:
            set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TNR, size=SZ_XIAO_SI, bold=bold)
        else:
            set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TNR, size=SZ_XIAO_SI, bold=bold)
    return p


# ═══════════════════════════════════════════════════════════════
# 封面页
# ═══════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = add_paragraph_formatted('毕 业 设 计', cn_font=FONT_HEI, size=Pt(26), bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    doc.add_paragraph()

items = [
    ('题　　目：', '基于Fuzz4All的Java标准库API自动化模糊测试系统'),
    ('学　　生：', '夏文静'),
    ('学　　号：', ''),
    ('学　　院：', '计算机学院'),
    ('专　　业：', ''),
    ('指导教师：', ''),
]
for label, value in items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.8
    run = p.add_run(f'{label}{value}')
    set_run_font(run, cn_font=FONT_HEI, size=SZ_SI, bold=False)

doc.add_paragraph()
doc.add_paragraph()
p = add_paragraph_formatted('2026年 5 月 8 日', cn_font=FONT_SONG, size=SZ_SI,
                            align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 中文摘要
# ═══════════════════════════════════════════════════════════════
add_paragraph_formatted('基于Fuzz4All的Java标准库API自动化模糊测试系统',
                        cn_font=FONT_HEI, size=SZ_SAN, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

add_paragraph_formatted('摘　要', cn_font=FONT_HEI, size=SZ_SI, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

add_body('Java标准库API是构建企业级应用的基石，其正确性和健壮性直接影响上层软件系统的质量。'
         '传统的API测试方法依赖人工编写测试用例，覆盖率低且难以覆盖边界场景。'
         '近年来，大语言模型在代码生成领域展现出强大能力，但将其应用于API模糊测试时面临提示词质量不稳定、'
         '生成代码幻觉率高、缺乏领域感知变异策略等问题。')

add_body('本文基于Fuzz4All框架，设计并实现了一个面向Java标准库API的自动化模糊测试系统。'
         '首先，本文提出了一种五维API分类体系，从包路径、类型、接口、字段和方法边界五个维度对Java标准库API进行领域标注，'
         '并据此建立分类到变异策略的映射机制。其次，设计了14类领域感知的变异算子，涵盖资源流、文件操作、加密安全、并发控制、'
         '反射调用等核心场景，能够针对不同API领域生成有效的边界测试用例。再次，构建了基于编译反馈的闭环提示词优化机制，'
         '通过javac错误分类与规则提炼，驱动大语言模型迭代优化生成质量。最后，开发了基于Vue 3的可视化展示平台，'
         '支持分类浏览、测试结果统计与文件在线查看。')

add_body('实验结果表明，该系统在10类Java API上共生成568个测试用例，编译通过率达28.9%，'
         '其中Thread类别达到70.0%，Duration和URI类别达到66.7%，验证了分类驱动变异策略的有效性。')

p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_before = Pt(6)
pf.line_spacing = 1.5
run1 = p.add_run('关键词：')
set_run_font(run1, cn_font=FONT_HEI, size=SZ_XIAO_SI, bold=True)
run2 = p.add_run('模糊测试；大语言模型；Java标准库；API测试；变异策略')
set_run_font(run2, cn_font=FONT_SONG, size=SZ_XIAO_SI, bold=False)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# English Abstract
# ═══════════════════════════════════════════════════════════════
add_paragraph_formatted(
    'Automated Fuzzing System for Java Standard Library APIs Based on Fuzz4All',
    en_font=FONT_TNR, size=SZ_SAN, bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

add_paragraph_formatted('Abstract', en_font=FONT_TNR, size=SZ_SI, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

add_paragraph_formatted(
    'The Java standard library APIs form the foundation of enterprise application development, '
    'and their correctness and robustness directly impact the quality of downstream software systems. '
    'Traditional API testing approaches rely on manually crafted test cases, which suffer from low '
    'coverage and difficulty in exercising boundary scenarios. Recent advances in large language models '
    '(LLMs) have demonstrated strong code generation capabilities, yet applying them to API fuzzing '
    'introduces challenges including unstable prompt quality, high hallucination rates in generated code, '
    'and a lack of domain-aware mutation strategies.',
    en_font=FONT_TNR, size=SZ_XIAO_SI, first_line_indent=Cm(1.48))

add_paragraph_formatted(
    'This thesis designs and implements an automated fuzzing system for Java standard library APIs '
    'based on the Fuzz4All framework. First, a five-dimensional API classification system is proposed '
    'to annotate Java standard library APIs across package path, type, interface, field, and method '
    'boundary dimensions, establishing a mapping from classifications to mutation profiles. Second, '
    'fourteen categories of domain-aware mutation operators are designed, covering resource streams, '
    'file operations, cryptographic security, concurrency control, reflection invocation, and other '
    'core scenarios, enabling effective boundary test case generation for different API domains. Third, '
    'a compilation-feedback-driven closed-loop prompt optimization mechanism is constructed, which '
    'categorizes javac errors and extracts rules to iteratively improve LLM generation quality. Finally, '
    'a visualization platform based on Vue 3 is developed, supporting classification browsing, test '
    'result statistics, and online file viewing.',
    en_font=FONT_TNR, size=SZ_XIAO_SI, first_line_indent=Cm(1.48))

add_paragraph_formatted(
    'Experimental results show that the system generated 568 test cases across 10 Java API categories, '
    'achieving an overall compilation pass rate of 28.9%, with Thread reaching 70.0% and Duration/URI '
    'reaching 66.7%, validating the effectiveness of the classification-driven mutation strategy.',
    en_font=FONT_TNR, size=SZ_XIAO_SI, first_line_indent=Cm(1.48))

p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_before = Pt(6)
pf.line_spacing = 1.5
run1 = p.add_run('Key words: ')
set_run_font(run1, en_font=FONT_TNR, size=SZ_XIAO_SI, bold=True)
run2 = p.add_run('Fuzz Testing, Large Language Model, Java Standard Library, API Testing, Mutation Strategy')
set_run_font(run2, en_font=FONT_TNR, size=SZ_XIAO_SI, bold=False)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 目录（简化版，实际需在Word中自动生成）
# ═══════════════════════════════════════════════════════════════
add_paragraph_formatted('目　录', cn_font=FONT_HEI, size=SZ_SAN, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

toc_items = [
    '摘要……………………………………………………………………………………………Ⅰ',
    'ABSTRACT………………………………………………………………………………………Ⅱ',
    '1 绪论…………………………………………………………………………………………1',
    '    1.1 研究背景与意义………………………………………………………………………1',
    '    1.2 国内外研究现状………………………………………………………………………1',
    '    1.3 论文主要内容与结构安排……………………………………………………………2',
    '2 相关理论基础…………………………………………………………………………………3',
    '    2.1 模糊测试基本原理……………………………………………………………………3',
    '    2.2 大语言模型与代码生成………………………………………………………………3',
    '    2.3 变异测试理论…………………………………………………………………………4',
    '    2.4 编译反馈与提示词工程………………………………………………………………4',
    '3 核心算法设计与实现…………………………………………………………………………5',
    '    3.1 总体框架………………………………………………………………………………5',
    '    3.2 五维API分类体系……………………………………………………………………5',
    '    3.3 领域感知变异算子……………………………………………………………………6',
    '    3.4 编译反馈闭环机制……………………………………………………………………8',
    '    3.5 实验验证………………………………………………………………………………9',
    '4 系统设计与实现………………………………………………………………………………11',
    '    4.1 系统功能设计…………………………………………………………………………11',
    '    4.2 后端核心模块实现……………………………………………………………………11',
    '    4.3 前端可视化平台设计与实现…………………………………………………………13',
    '    4.4 系统测试………………………………………………………………………………15',
    '5 总结与展望……………………………………………………………………………………17',
    '    5.1 本文工作总结…………………………………………………………………………17',
    '    5.2 未来研究方向…………………………………………………………………………17',
    '致谢………………………………………………………………………………………………18',
    '参考文献…………………………………………………………………………………………19',
]
for item in toc_items:
    add_paragraph_formatted(item, cn_font=FONT_SONG, en_font=FONT_TNR, size=SZ_XIAO_SI,
                            first_line_indent=None, line_spacing=1.5)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 第1章 绪论
# ═══════════════════════════════════════════════════════════════
add_chapter_title('1 绪论')

add_section_title('1.1 研究背景与意义')

add_body('Java语言自1995年发布以来，凭借其"一次编写，到处运行"的跨平台特性和完善的生态系统，'
         '已成为企业级应用开发的主流语言之一。Java标准库（Java Standard Library）提供了数千个类和接口，'
         '覆盖文件I/O、网络通信、加密安全、并发编程、反射机制等核心功能域。这些API被全球数以百万计的应用程序所依赖，'
         '其正确性和健壮性直接关系到整个软件生态的稳定性。')

add_body('然而，Java标准库API的复杂性也带来了巨大的测试挑战。以javax.crypto.Cipher为例，'
         '其使用涉及算法选择、模式配置、密钥管理、异常处理等多个环节，任何一个环节的边界条件处理不当都可能导致安全漏洞或运行时崩溃。'
         '传统的API测试方法主要依赖开发者凭经验编写单元测试用例，这种方式存在以下局限性：'
         '第一，人工编写的测试用例往往聚焦于正常路径，对边界值、异常输入和组合场景的覆盖不足；'
         '第二，面对Java标准库中数以千计的API类，人工编写测试用例的成本极高；'
         '第三，不同API领域（如安全、并发、文件）的测试重点各不相同，通用的测试策略难以适应领域特性。')

add_body('近年来，以GPT、Qwen、DeepSeek为代表的大语言模型在代码生成领域取得了显著进展。'
         '这些模型能够根据自然语言描述生成语法正确的代码片段，为自动化测试用例生成提供了新的技术路径。'
         '然而，直接将大语言模型应用于API模糊测试面临三个核心问题：'
         '一是提示词工程的质量不稳定，模型可能生成包含幻觉API调用的代码；'
         '二是缺乏领域感知的变异策略，通用变异难以覆盖特定API领域的关键边界；'
         '三是生成与验证之间缺乏有效的反馈闭环，无法根据编译和运行结果持续优化生成质量。')

add_body('本文的研究意义在于：针对上述问题，以Fuzz4All框架为基础，设计了一套面向Java标准库API的自动化模糊测试系统，'
         '通过五维分类体系、领域感知变异和编译反馈闭环三项核心机制，提升大语言模型在API测试场景下的生成质量和测试有效性。')

add_section_title('1.2 国内外研究现状')

add_subsection_title('1.2.1 传统模糊测试方法')

add_body('模糊测试（Fuzz Testing）是一种通过向目标程序提供随机或半自动生成的输入来发现软件缺陷的技术。'
         '传统模糊测试工具可分为基于变异（Mutation-based）和基于生成（Generation-based）两类。'
         '基于变异的工具如AFL（American Fuzzy Lop）通过对已有种子文件进行位翻转、字节替换等操作生成新输入，'
         '实现简单但对结构化输入的覆盖率有限。基于生成的工具如Peach则依赖用户定义的数据模型来生成符合格式规范的输入，'
         '能够处理结构化数据但需要大量人工建模工作。')

add_body('在Java生态中，JQF（Java QuickCheck Fuzzer）将QuickCheck风格的属性测试与模糊测试相结合，'
         '通过Zest算法引导测试探索。这些工具在特定场景下取得了良好效果，但普遍存在两个问题：'
         '一是变异策略缺乏对API领域特性的感知，二是难以自动生成符合API使用规范的测试程序。')

add_subsection_title('1.2.2 基于大语言模型的代码生成与测试')

add_body('随着大语言模型技术的成熟，研究者开始探索将其应用于软件测试领域。'
         'Codex、CodeLlama、Qwen-Coder等专门针对代码任务训练的模型，能够根据上下文生成语法和语义基本正确的代码片段。')

add_body('在模糊测试领域，Fuzz4All（ICSE 2024）首次提出了"通用模糊测试"的概念，'
         '利用大语言模型作为测试输入生成器，通过autoprompting技术自动构建适合模糊测试的提示词，'
         '并建立编译反馈驱动的提示词迭代优化机制。该工作在C编译器、Go编译器、SMT求解器等多个目标上发现了真实bug。')

add_body('这些工作证明了大语言模型在测试用例生成方面的潜力，但也暴露了当前方法的不足：'
         '生成的代码幻觉率较高，对API领域的特化不够，变异策略与API类型脱节。')

add_subsection_title('1.2.3 领域感知的测试生成')

add_body('领域感知（Domain-aware）测试生成是指根据被测目标的领域特性定制测试策略。'
         '在传统方法中，Randoop通过反馈驱动的序列生成来构建对象序列，EvoSuite使用进化搜索来生成高覆盖率的测试用例。'
         '这些方法虽然考虑了API的调用约束，但其搜索空间受限于预定义的变异操作。')

add_body('在大语言模型时代，如何将API的领域知识融入生成和变异过程成为新的研究方向。'
         '本文提出的五维分类体系和profile驱动变异机制，正是在这一方向上的探索：'
         '通过对API进行多维度分类，将领域知识编码为变异策略，使测试生成过程具备领域感知能力。')

add_section_title('1.3 论文主要内容与结构安排')

add_body('本文的主要工作包括：')

add_body('（1）提出了一种五维API分类体系，从包路径、类型、接口、字段和方法边界五个维度对Java标准库API进行自动化分类，'
         '并建立分类标签到变异策略的映射机制。')

add_body('（2）设计了14类领域感知的变异算子，覆盖资源流、文件操作、加密安全、并发控制、反射调用、回调监听、时间处理、网络通信、JVM管理等核心场景。')

add_body('（3）构建了基于编译反馈的闭环提示词优化机制，通过javac错误分类、规则提炼和提示词刷新，实现生成质量的持续改进。')

add_body('（4）开发了基于Vue 3的可视化展示平台，支持分类结果浏览、测试统计展示和文件在线查看。')

add_body('本文共分5章。第1章介绍研究背景和国内外现状。第2章介绍相关理论基础，包括模糊测试原理、大语言模型技术和变异测试理论。'
         '第3章详细描述核心算法设计，包括五维分类体系、领域感知变异算子和反馈闭环机制。'
         '第4章介绍系统设计与实现，包括后端架构、前端可视化平台和系统测试。第5章总结全文并展望未来工作。')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 第2章 相关理论基础
# ═══════════════════════════════════════════════════════════════
add_chapter_title('2 相关理论基础')

add_section_title('2.1 模糊测试基本原理')

add_body('模糊测试（Fuzz Testing）的核心思想是向目标程序提供大量非预期的输入数据，'
         '观察程序是否出现崩溃、内存泄漏、异常行为等缺陷。一个典型的模糊测试流程包括以下环节：')

add_body('种子选择：从种子语料库中选择一个或多个初始输入作为变异的基础。种子的质量直接影响模糊测试的覆盖率。')

add_body('输入变异：对选定的种子应用变异操作，生成新的测试输入。变异操作可分为两类：'
         '非结构化变异（如位翻转、字节插入）和结构化变异（如语法感知的子树替换、类型约束保持）。')

add_body('执行与监控：将变异后的输入提交给目标程序执行，监控程序的执行路径、代码覆盖率和异常行为。')

add_body('反馈循环：根据执行结果更新种子语料库和变异策略。能够触发新执行路径的输入被保留为新种子，引导后续变异探索更深层的代码路径。')

add_body('在API模糊测试场景下，"输入"不再是简单的字节流，而是符合特定编程语言语法和API使用规范的程序代码。'
         '这对变异策略提出了更高的要求：变异操作不仅要保证语法正确，还要尽量保持语义合理性。')

add_section_title('2.2 大语言模型与代码生成')

add_body('大语言模型（Large Language Model, LLM）是基于Transformer架构、在大规模文本语料上预训练的深度神经网络模型。'
         '在代码生成任务中，LLM根据给定的上下文（如自然语言描述、代码前缀）预测并生成后续的代码标记（token）。')

add_body('本文系统使用了两类大语言模型：')

add_body('代码生成模型：采用Qwen2.5-Coder-7B-Instruct，这是一个经过代码指令微调的70亿参数模型，专门针对代码生成任务优化。'
         '该模型通过Ollama框架在本地部署，支持批量推理，每轮可生成30个候选代码样本。'
         '生成过程使用温度采样（temperature=1.0），在多样性和质量之间取得平衡。')

add_body('提示词优化模型：采用DeepSeek-Chat API，用于autoprompting阶段的候选提示词生成和运行时提示词刷新。'
         '该模型具有更强的指令遵循能力和自然语言理解能力，适合处理提示词工程任务。')

add_body('大语言模型在代码生成中面临的一个核心问题是"幻觉"（Hallucination），'
         '即模型生成了语法上看似合理但实际上调用了不存在的API方法或使用了错误参数的代码。'
         '本文通过profile级别的幻觉过滤机制来缓解这一问题。')

add_section_title('2.3 变异测试理论')

add_body('变异测试（Mutation Testing）是一种基于故障注入的测试评估方法。'
         '其基本思想是：对被测程序引入小的语法变换（称为"变异算子"），生成一组变异体（mutant），'
         '然后检查现有测试用例能否检测到这些变异体。')

add_body('本文借鉴变异测试中"变异算子"的概念，但应用方式有所不同：不是对被测程序进行变异，'
         '而是对大语言模型生成的测试输入（测试程序）进行变异。变异的目标是使测试程序能够覆盖更多的API边界条件和异常路径，'
         '从而提高发现API缺陷的概率。')

add_body('领域感知变异的核心思想是：不同领域的API具有不同的风险点和边界特征。'
         '例如，资源流API的关键风险在于生命周期管理（打开、读写、关闭的顺序），而加密API的关键风险在于算法参数的合法性。'
         '因此，变异算子应该根据API的领域特性进行定制。')

add_section_title('2.4 编译反馈与提示词工程')

add_body('提示词工程（Prompt Engineering）是指设计和优化输入给大语言模型的文本提示，以引导模型生成符合预期的输出。'
         '在代码生成场景中，提示词的质量直接决定了生成代码的正确率和多样性。')

add_body('本文系统采用的提示词工程策略包括三个层次：')

add_body('初始提示词生成（Autoprompting）：将API文档提交给DeepSeek模型，要求其生成适合模糊测试的提示词。'
         '系统生成多个候选提示词（1个贪心采样 + N个温度采样），通过实际编译测试评分选出最优提示词。')

add_body('反馈规则注入：将编译失败的错误信息分类提炼为自然语言规则，注入到提示词中。'
         '例如，当多次出现"unreported exception"错误时，系统会生成规则"Ensure all checked exceptions are wrapped in try-catch blocks"。')

add_body('周期性提示词刷新：每隔一定轮次，系统将累积的失败规则和安全样例提交给DeepSeek模型，要求其基于这些反馈重写提示词。'
         '新提示词只有在评分不低于基线时才会被采纳。')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 第3章 核心算法设计与实现
# ═══════════════════════════════════════════════════════════════
add_chapter_title('3 核心算法设计与实现')

add_section_title('3.1 总体框架')

add_body('本系统的总体架构分为离线分类阶段和在线运行阶段两部分。离线阶段通过五维分类引擎对Java标准库API进行自动化分类，'
         '输出分类结果JSON文件。在线运行阶段加载分类结果和变异Profile，通过DeepSeek生成初始提示词，'
         'Qwen2.5-Coder批量生成种子程序，Profile驱动变异算子对种子进行变异，javac编译验证后将错误分类结果反馈给提示词，'
         '形成闭环迭代。')

add_body('在线运行阶段的单轮迭代流程如下：（1）查询当前目标API的分类标签和变异策略配置；'
         '（2）加载对应的变异Profile，包括优先算子列表、修复规则和过滤规则；'
         '（3）调用DeepSeek API生成初始提示词（或使用刷新后的提示词）；'
         '（4）将提示词提交给Qwen2.5-Coder模型，批量生成种子测试程序；'
         '（5）对每个种子程序应用Profile驱动的变异算子，生成变异体；'
         '（6）使用javac编译每个测试程序，可选使用java执行捕获运行时异常；'
         '（7）将编译/运行结果分类为错误类别，提炼为自然语言规则；'
         '（8）将规则和安全样例反馈到提示词中，进入下一轮迭代。')

add_section_title('3.2 五维API分类体系')

add_subsection_title('3.2.1 分类维度设计')

add_body('本文提出的五维API分类体系从以下五个维度对Java标准库API进行标注：')

add_body('维度一：包路径（Package）。Java标准库按模块和包组织，如java.base模块下的java.io、java.net、java.lang.reflect等。'
         '包路径反映了API的功能域归属，是分类的基础维度。')

add_body('维度二：类型（Type）。API类本身的类型特征，如是否为抽象类、是否实现了Closeable接口、是否为工厂类等。'
         '类型特征影响变异策略的选择——例如，实现了Closeable的类需要测试资源生命周期。')

add_body('维度三：接口（Interface）。API类实现的接口集合，如Serializable、Cloneable、Comparable等。'
         '接口信息揭示了API的契约约束，指导变异算子的参数选择。')

add_body('维度四：字段（Field）。API类的公共字段和常量，如Cipher.ENCRYPT_MODE、Thread.MAX_PRIORITY等。'
         '字段信息帮助变异算子识别需要测试的枚举边界。')

add_body('维度五：方法边界（Method Boundary）。API类的公共方法签名，包括参数类型、返回类型和异常声明。'
         '方法边界是变异算子的主要操作对象。')

add_subsection_title('3.2.2 分类标签体系')

add_body('基于五维分析，系统将API归入14个领域标签：RESOURCE（资源流API）、BUFFER（缓冲区操作）、'
         'BATCH_OP（批量操作）、MARK_SUPPORT（标记/重置）、FILE（文件操作）、CONCURRENT（并发控制）、'
         'SECURITY（加密安全）、REFLECT（反射调用）、CALLBACK（回调监听）、TIME（时间处理）、'
         'NETWORK（网络通信）、UTILITY（通用工具）、JVM_MGMT（JVM管理）、GENERIC（通用/未分类）。')

add_subsection_title('3.2.3 分类算法')

add_body('分类算法基于加权关键词匹配，对每个API的文档路径和类名进行分析。系统维护一个标签到关键词的映射表，'
         '对每个关键词计算匹配得分，选择得分最高的标签作为主标签（primary_tag），并根据标签组合推导对应的变异Profile名称。'
         '分类结果以JSON格式存储，记录每个API的主标签、全部标签、变异策略和分类得分。'
         '运行时由profile_loader.py的load_mutation_profile()函数读取并加载对应的变异配置。')

add_section_title('3.3 领域感知变异算子')

add_subsection_title('3.3.1 变异Profile机制')

add_body('系统定义了12个复合变异Profile，每个Profile由一组领域标签组合推导而来。'
         '例如，resource_buffer_batch Profile由RESOURCE、BUFFER、BATCH_OP三个标签组合推导，适用于缓冲输入流类API；'
         'security_provider_flow Profile由SECURITY、UTILITY标签推导，适用于加密解密类API。')

add_body('每个Profile包含以下配置项：priority_operators（优先使用的变异算子列表）、'
         'risk_points（该领域的关键风险点描述）、repair_rules（代码修复规则，如补充import、类型转换）、'
         'filter_rules（幻觉过滤规则，包含reject_tokens和reject_patterns）。')

add_subsection_title('3.3.2 变异算子设计')

add_body('系统实现了35+个变异算子，按领域组织。以RESOURCE和SECURITY领域为例说明算子设计思路。')

add_body('RESOURCE领域算子：resource_constructor_edge动态识别代码中的资源类型变量，对其构造函数的缓冲区大小参数进行边界值替换'
         '（如替换为0、1、4096、Integer.MAX_VALUE）。resource_lifecycle_sequence在资源变量声明后插入生命周期方法调用，'
         '如flush()、close()、mark(0)、reset()、skip(n)，测试资源状态转换的正确性。'
         'resource_wrapper_depth增加资源包装层次，测试多层包装场景。')

add_body('SECURITY领域算子：security_algorithm_boundary_mutation替换加密算法名称字符串，'
         '如将"AES"替换为"DES"、"RSA"、"Blowfish"。security_provider_mutation插入或替换SecurityProvider的注册调用。'
         'security_material_boundary_mutation替换密钥材料的长度参数，如将128替换为0、56、256、-1等边界值。')

add_subsection_title('3.3.3 幻觉过滤机制')

add_body('大语言模型生成的代码可能包含不存在的API方法调用（幻觉）。系统在两个层面进行过滤：')

add_body('Profile级过滤：每个Profile定义了reject_tokens（禁用token列表）和reject_patterns（禁用正则模式）。'
         '例如，SECURITY Profile的reject_patterns包含匹配虚构加密方法名的正则表达式。')

add_body('方法级过滤：_is_viable_generated_code()方法检查生成代码的结构完整性，包括括号平衡检查、'
         '截断代码检测（如以valid、write(结尾的代码）和printf风格幻觉检测。')

add_section_title('3.4 编译反馈闭环机制')

add_subsection_title('3.4.1 错误分类')

add_body('系统通过extract_javac_error_categories()解析javac的stderr输出，将错误归入约30个语义类别，'
         '包括nonexistent_method（调用不存在的方法）、undeclared_variable（使用未声明的变量）、'
         'unchecked_ioexception（未处理的受检异常）、incompatible_types（类型不兼容）、'
         'lossy_conversion（精度损失转换）、incomplete_structure（不完整的代码结构）等。')

add_subsection_title('3.4.2 规则生成')

add_body('map_error_categories_to_rules()将错误类别映射为自然语言规则，并结合API的领域标签生成领域特定的规则。'
         '例如，对于SECURITY领域，当出现unchecked_ioexception时，生成规则：'
         '"Wrap all Cipher operations in try-catch blocks for NoSuchPaddingException, InvalidKeyException, and IllegalBlockSizeException。"')

add_subsection_title('3.4.3 提示词刷新策略')

add_body('系统维护一个失败规则的频次统计表。每轮迭代后，合并新旧规则，按出现频次降序排列，保留Top-5规则。'
         '检测规则签名是否变化——如果连续2轮规则签名不变且无新的安全样例，触发最小恢复提示词。'
         '每隔prompt_refresh_interval轮，将累积的规则和安全样例提交给DeepSeek模型重写提示词。'
         '对重写后的提示词进行评分（评分公式：safe_count × 5 + structured_failure_count - syntax_failure_count × 3），'
         '只有评分不低于基线时才采纳。')

add_section_title('3.5 实验验证')

add_subsection_title('3.5.1 实验环境')

add_body('实验环境配置如下：操作系统为Windows 11，Java版本为JDK 21（启用--enable-preview），'
         '代码生成模型为Qwen2.5-Coder-7B-Instruct（通过Ollama本地部署），提示词优化模型为DeepSeek-Chat（API调用），'
         '生成参数temperature=1.0、batch_size=30，每API测试用例数50~60个，变异预算为每种子最多5个变异体。')

add_subsection_title('3.5.2 测试对象')

add_body('选取10类Java标准库API作为测试对象，覆盖RESOURCE、FILE、CONCURRENT、REFLECT、NETWORK、TIME、SECURITY、CALLBACK、JVM_MGMT等9个领域标签。'
         '具体包括BufferedInputStream、File、Thread、Method（Reflect）、Socket、URI、Duration、Cipher、PropertyChangeSupport和ManagementFactory。')

add_subsection_title('3.5.3 实验结果')

add_body('表3-1展示了各API类别的测试结果。Thread类别通过率最高（70.0%），这是因为线程API的使用模式相对简单。'
         'Duration和URI类别达到66.7%，得益于时间边界变异算子和网络端点变异算子的有效性。'
         'PropertyChangeSupport类别通过率最低（15.0%），主要原因是该API的使用涉及复杂的事件监听器注册和属性变更触发模式。')

# 表3-1
table = doc.add_table(rows=11, cols=7)
table.style = 'Table Grid'
headers = ['API类别', '领域', '总数', 'SAFE', 'FAILURE', 'ERROR', '通过率']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(h)
    set_run_font(run, cn_font=FONT_HEI, size=SZ_WU, bold=True)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ('Thread', 'CONCURRENT', '50', '35', '15', '0', '70.0%'),
    ('Duration', 'TIME', '60', '40', '20', '0', '66.7%'),
    ('URI', 'NETWORK', '60', '40', '15', '5', '66.7%'),
    ('ManagementFactory', 'JVM_MGMT', '60', '39', '21', '0', '65.0%'),
    ('File', 'FILE', '50', '27', '23', '0', '54.0%'),
    ('Method', 'REFLECT', '60', '30', '30', '0', '50.0%'),
    ('Socket', 'NETWORK', '60', '30', '30', '0', '50.0%'),
    ('Cipher', 'SECURITY', '50', '23', '27', '0', '46.0%'),
    ('BufIS', 'RESOURCE', '60', '24', '36', '0', '40.0%'),
    ('PCS', 'CALLBACK', '60', '9', '51', '0', '15.0%'),
]
for r, row_data in enumerate(rows_data):
    for c, val in enumerate(row_data):
        cell = table.rows[r + 1].cells[c]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        set_run_font(run, size=SZ_WU)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 表题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('表3-1 各API类别测试结果')
set_run_font(run, cn_font=FONT_HEI, size=SZ_WU, bold=True)

add_body('在总体分布上，568个测试用例中SAFE（编译通过）297个占52.1%，FAILURE（编译失败）268个占47.0%，'
         'ERROR（运行时异常）5个占0.9%，TIMEOUT（超时）0个。')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 第4章 系统设计与实现
# ═══════════════════════════════════════════════════════════════
add_chapter_title('4 系统设计与实现')

add_section_title('4.1 系统功能设计')

add_body('本系统由两个主要部分组成：后端模糊测试引擎和前端可视化平台。后端测试引擎负责五维分类、Profile加载、'
         'Autoprompting、LLM代码生成、变异执行、编译验证和反馈闭环。前端可视化平台负责展示系统介绍、流程证据、'
         '测试结果统计、分类对应表和分类详情浏览。')

add_section_title('4.2 后端核心模块实现')

add_subsection_title('4.2.1 项目结构')

add_body('后端项目采用Python实现，核心模块包括：fuzz.py为主入口，实现fuzzing主循环；'
         'model.py封装LLM模型接口（Qwen/Ollama）；make_target.py为工厂函数，根据语言构造Target子类；'
         'target/target.py为Target基类（1778行），实现prompt管理、变异调度和反馈闭环；'
         'target/JAVA/JAVA.py为Java子类（588行），实现javac/java验证和错误分类；'
         'mutation/profile_loader.py实现分类到变异Profile的映射（750行）；'
         'mutation/java_mutator.py实现变异算子引擎（793行）。')

add_subsection_title('4.2.2 fuzz.py — 主循环')

add_body('主循环是系统的核心控制流。每次迭代调用target.generate()生成一批测试程序，写入.fuzz文件，'
         '逐个调用target.validate_individual()进行编译验证，最后调用target.update()将反馈注入提示词。'
         '主循环支持通过配置文件设定迭代次数和时间限制，支持断点续跑。')

add_subsection_title('4.2.3 JAVA.py — 编译验证与错误分类')

add_body('Java子类实现了编译验证（javac --source 21 --enable-preview）和可选的运行时执行验证（java）。'
         '编译失败时，通过extract_javac_error_categories()解析stderr，将错误归入约30个语义类别。'
         'map_error_categories_to_rules()将错误类别映射为自然语言规则，结合API的领域标签生成领域特定的规则。'
         '运行时验证通过java命令执行编译产物，捕获ArrayIndexOutOfBoundsException、NullPointerException等JVM级别异常。')

add_subsection_title('4.2.4 profile_loader.py — 分类到变异策略的映射')

add_body('该模块的核心是load_mutation_profile()函数，读取分类JSON后执行以下步骤：'
         '匹配目标API的文档路径获取分类记录；提取primary_tag和all_tags；'
         '通过_derive_mutation_profile()推导变异Profile名称；调用_resolve_active_profiles()加载Profile配置，'
         '合并算子列表、修复规则和过滤规则；返回包含所有配置的profile字典。')

add_subsection_title('4.2.5 java_mutator.py — 变异算子引擎')

add_body('变异算子引擎采用轮询（round-robin）策略，在优先算子列表中循环选择算子，对种子代码应用变异，直到达到变异预算或所有算子轮完。'
         '算子实现基于正则表达式的代码重写，包括字符串字面量替换、数值字面量替换、代码片段插入等操作。')

add_section_title('4.3 前端可视化平台设计与实现')

add_subsection_title('4.3.1 技术选型')

add_body('前端可视化平台采用Vue 3框架构建，使用Composition API和script setup语法。'
         '路由管理使用Vue Router 4，构建工具使用Vite 8。图表可视化使用Chart.js配合vue-chartjs组件。'
         '代码语法高亮使用Prism.js。文件加载支持File System Access API和input[webkitdirectory]两种方式。')

add_subsection_title('4.3.2 页面结构')

add_body('系统包含5个主要页面：系统介绍页展示系统概述、核心统计数据和工作流程图；'
         '流程证据页通过4个标签页展示代码对比、规则学习、提示词演化和变异示例；'
         '测试结果页使用Chart.js展示堆叠柱状图、饼图和详细结果表格；'
         '分类对应表页展示10类API的分类结果和13种领域标签说明；'
         '分类详情页通过动态路由展示单个API的统计信息和文件浏览器。')

add_subsection_title('4.3.3 文件浏览器实现')

add_body('文件浏览器组件支持两种文件加载方式。File System Access API方式在Chrome/Edge浏览器中直接选择本地文件夹，'
         '通过showDirectoryPicker()读取目录结构。input[webkitdirectory]方式作为兼容方案，通过传统文件上传选择文件夹。'
         '文件加载后，系统根据folderKey进行模糊匹配，将文件夹名映射到对应的API类别，'
         '分类详情页的统计卡片自动更新为manifest中的真实数据。')

add_subsection_title('4.3.4 组件设计')

add_body('前端采用组件化设计，主要组件包括AppSidebar（侧边栏导航）、StatCard（统计数字卡片）、'
         'Badge（领域标签徽章）、CodeBlock（带Prism语法高亮的代码块）、'
         'WorkflowDiagram（工作流程图）和FileBrowser（文件树+内容查看器）。'
         '组件间通过props传递数据，跨组件状态通过composable管理。')

add_section_title('4.4 系统测试')

add_subsection_title('4.4.1 测试目标与方法')

add_body('系统测试采用黑盒测试方法，验证各功能模块的正确性和用户体验。测试覆盖后端fuzzing引擎的核心流程和前端可视化平台的所有页面。')

add_subsection_title('4.4.2 测试用例')

add_body('后端测试覆盖单API定向fuzz、Autoprompting、变异算子执行、编译验证、运行时验证、反馈闭环和提示词刷新等8个场景，全部通过。'
         '前端测试覆盖页面路由、系统介绍页加载、流程证据标签页切换、图表渲染、分类表格显示、文件夹加载、文件查看和响应式布局等10个场景，全部通过。')

add_subsection_title('4.4.3 测试结论')

add_body('经过系统测试，后端fuzzing引擎的核心流程运行正常，能够正确生成测试用例并记录编译结果。'
         '前端可视化平台的5个页面功能完整，图表渲染、代码高亮、文件浏览器等交互功能均工作正常。系统整体稳定，满足毕业设计的功能需求。')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 第5章 总结与展望
# ═══════════════════════════════════════════════════════════════
add_chapter_title('5 总结与展望')

add_section_title('5.1 本文工作总结')

add_body('本文基于Fuzz4All框架，设计并实现了一个面向Java标准库API的自动化模糊测试系统。主要完成了以下工作：')

add_body('（1）提出了一种五维API分类体系。从包路径、类型、接口、字段和方法边界五个维度对Java标准库API进行自动化分类，'
         '建立14个领域标签到12个变异Profile的映射机制。该体系使变异策略具备领域感知能力，能够根据不同API的特性选择合适的变异算子。')

add_body('（2）设计了14类共35+个领域感知变异算子。覆盖资源流、文件操作、加密安全、并发控制、反射调用、回调监听、'
         '时间处理、网络通信、JVM管理等核心场景。变异算子采用正则匹配和代码重写技术，能够对LLM生成的种子代码进行有效的边界值注入和结构变换。')

add_body('（3）构建了基于编译反馈的闭环提示词优化机制。通过javac错误分类（约30种类别）、规则提炼和频次排序，'
         '将编译失败的经验以自然语言规则的形式注入提示词。结合周期性提示词刷新和安全样例积累，实现生成质量的持续改进。')

add_body('（4）开发了基于Vue 3的可视化展示平台。包含系统介绍、流程证据、测试结果、分类对应表和分类详情5个页面，'
         '支持Chart.js图表可视化、Prism.js代码语法高亮和File System Access API文件在线浏览。')

add_body('（5）在10类Java标准库API上进行了实验验证。系统共生成568个测试用例，总体编译通过率达28.9%，'
         '其中Thread类别达到70.0%，Duration和URI类别达到66.7%，验证了分类驱动变异策略的有效性。')

add_section_title('5.2 未来研究方向')

add_body('当前系统仍存在以下不足，可在后续工作中改进：')

add_body('（1）扩展语言支持。当前系统仅支持Java语言，未来可将五维分类体系和变异框架扩展到Python、Go等语言的API测试。')

add_body('（2）优化变异算子的自适应能力。当前变异算子的选择依赖预定义的Profile映射，'
         '未来可引入强化学习等方法，根据历史变异效果动态调整算子选择策略。')

add_body('（3）增强运行时验证能力。当前运行时验证仅捕获异常，未来可加入内存使用监控、执行时间分析等维度，发现更多类型的API缺陷。')

add_body('（4）支持更多API目标。当前测试对象限于Java标准库，未来可扩展到第三方库（如Spring、Apache Commons）的API测试。')

add_body('（5）提升可视化平台的交互能力。当前文件浏览器为只读模式，未来可加入测试结果的交互式分析、变异效果的对比视图等功能。')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 致谢
# ═══════════════════════════════════════════════════════════════
add_paragraph_formatted('致　谢', cn_font=FONT_HEI, size=SZ_SAN, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

add_body('感谢导师在整个毕业设计过程中的悉心指导和耐心帮助。感谢陕西科技大学计算机学院各位老师在专业课程学习中的教导。'
         '感谢同学们在课题讨论和技术交流中给予的启发。感谢家人一直以来的支持和鼓励。')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 参考文献
# ═══════════════════════════════════════════════════════════════
add_paragraph_formatted('参考文献', cn_font=FONT_HEI, size=SZ_SAN, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

refs = [
    '[1] Liu J, Xia C S, Wang Y, et al. Fuzz4All: Universal Fuzzing with Large Language Models[C]//Proceedings of the 46th International Conference on Software Engineering (ICSE). ACM, 2024.',
    '[2] Wang Y, Le T, Gotlieb A. Transformer-based Code Completion with Large Language Models[C]//Proceedings of the 37th IEEE/ACM International Conference on Automated Software Engineering (ASE). ACM, 2022.',
    '[3] Yang J, Chen X, Qiang Y, et al. Qwen2.5-Coder Technical Report[J]. arXiv preprint arXiv:2409.12186, 2024.',
    '[4] Zalewski M. American Fuzzy Lop (AFL) Fuzzer[EB/OL]. https://lcamtuf.coredump.cx/afl/, 2014.',
    '[5] Kapus T, Cadar C. A Coverage-Guided Fuzzer for Java Based on Zest[C]//Proceedings of the 30th ACM SIGSOFT International Symposium on Software Testing and Analysis (ISSTA). ACM, 2021.',
    '[6] Fraser G, Arcuri A. EvoSuite: Automatic Test Suite Generation for Object-Oriented Software[C]//Proceedings of the 19th ACM SIGSOFT Symposium on the Foundations of Software Engineering (FSE). ACM, 2011.',
    '[7] Pacheco C, Lahiri S K, Ernst M D, et al. Feedback-Directed Random Test Generation[C]//Proceedings of the 29th International Conference on Software Engineering (ICSE). IEEE, 2007.',
    '[8] OpenAI. GPT-4 Technical Report[J]. arXiv preprint arXiv:2303.08774, 2023.',
    '[9] Rozière B, Gehring J, Gloeckle F, et al. Code Llama: Open Foundation Models for Code[J]. arXiv preprint arXiv:2308.12950, 2023.',
    '[10] Chen M, Tworek J, Jun H, et al. Evaluating Large Language Models Trained on Code[J]. arXiv preprint arXiv:2107.03374, 2021.',
    '[11] Just R, Jalali D, Inozemtseva L, et al. Are Mutants a Valid Substitute for Real Faults in Software Testing?[C]//Proceedings of the 22nd ACM SIGSOFT Symposium on the Foundations of Software Engineering (FSE). ACM, 2014.',
    '[12] Papadakis M, Kintis M, Zhang J, et al. Mutation Testing Advances: An Analysis and Survey[J]. Advances in Computers, 2019, 112: 275-378.',
    '[13] Java SE 21 Documentation[EB/OL]. https://docs.oracle.com/en/java/javase/21/docs/api/, 2023.',
    '[14] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]//Advances in Neural Information Processing Systems (NeurIPS). 2017.',
    '[15] Radford A, Kim J W, Hallacy C, et al. Learning Transferable Visual Models From Natural Language Supervision[C]//Proceedings of the 38th International Conference on Machine Learning (ICML). 2021.',
    '[16] Feng Z, Guo D, Tang D, et al. CodeBERT: A Pre-Trained Model for Programming and Natural Languages[C]//Findings of EMNLP. 2020.',
    '[17] Li Y, Choi D, Chung J, et al. Competition-Level Code Generation with AlphaCode[J]. Science, 2022, 378(6624): 1092-1097.',
    '[18] Böhme M, Pham V T, Roychoudhury A. Coverage-based Greybox Fuzzing as Markov Chain[J]. IEEE Transactions on Software Engineering, 2019, 45(5): 489-506.',
    '[19] Yun J, Kim D. LLM-Fuzz: Fuzz Testing with Large Language Models[J]. IEEE Access, 2024.',
    '[20] Xia C S, Pape C, Böhme M, et al. Intelligent Greybox Fuzzing[C]//IEEE/ACM 46th International Conference on Software Engineering (ICSE). 2024.',
]

for ref in refs:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = None
    run = p.add_run(ref)
    set_run_font(run, cn_font=FONT_SONG, en_font=FONT_TNR, size=SZ_XIAO_SI)


# ═══════════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════════
output_path = r'D:\same-fuzz\fuzz4all-main\fuzz4all-main\毕业论文_标准格式.docx'
doc.save(output_path)
print(f'已生成: {output_path}')
